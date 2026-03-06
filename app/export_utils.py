from docx import Document
from io import BytesIO
from .services import db

def generate_project_docx(user_email, project_id):
    """
    從 Firestore 讀取專案資料，並生成 .docx 檔案流。
    """
    # 1. Fetch Project Metadata
    project_ref = db.collection('users').document(user_email).collection('projects').document(project_id)
    project_doc = project_ref.get()
    
    if not project_doc.exists:
        return None, "Project not found"
        
    project_data = project_doc.to_dict()
    report_title = project_data.get('report_title', 'Content Analysis Report')
    
    # 2. Fetch Pages
    pages_ref = project_ref.collection('pages').stream()
    pages = [p.to_dict() for p in pages_ref]
    
    # 3. Create Document
    doc = Document()
    doc.add_heading(report_title, 0)
    
    doc.add_paragraph(f"Generate Date: {project_data.get('created_at', 'N/A')}")
    doc.add_paragraph(f"Total Pages: {len(pages)}")
    doc.add_paragraph("-" * 50)
    
    for page in pages:
        url = page.get('url', 'Unknown URL')
        title = page.get('title', 'No Title')
        content = page.get('content', '(No Content)')
        status = page.get('status', 'Unknown')
        error = page.get('error')
        
        doc.add_heading(title, level=1)
        doc.add_paragraph(f"URL: {url}", style='Intense Quote')
        
        if status == 'success':
            doc.add_paragraph(content)
        else:
            doc.add_paragraph(f"[Error] Failed to crawl: {error}", style='List Bullet')
            
        doc.add_page_break()
        
    # 4. Save to BytesIO
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    filename = f"{report_title.replace(' ', '_')}.docx"
    return file_stream, filename
