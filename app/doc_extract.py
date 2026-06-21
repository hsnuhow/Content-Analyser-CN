# -*- coding: utf-8 -*-
"""上傳檔文字抽取層（自 project_routes.py 抽出）。

單一職責：把上傳檔 blob 轉成純文字。支援 .txt / .md / .docx（.doc 不支援）。
.txt/.md 純 stdlib；.docx 走 python-docx（lazy import）。無 db / Flask 依賴，可單獨測。
"""


def _decode_text(blob: bytes) -> str:
    """純文字解碼：先嚴格 UTF-8，失敗才退回 cp950（Big5，台灣 .txt 常見編碼），
    最後寬鬆 UTF-8 盡量保留。
    為何需要：台灣使用者上傳的純文字檔不少是 Big5/CP950，直接 utf-8 解會整段亂碼/掉字。
    順序很重要——UTF-8 先用嚴格模式（遇 Big5 位元組會丟 UnicodeDecodeError）才會落到 cp950；
    現今多數檔是 UTF-8，故 UTF-8 優先、Big5 當回退。"""
    for enc in ('utf-8', 'cp950'):
        try:
            return blob.decode(enc)
        except UnicodeDecodeError:
            continue
    return blob.decode('utf-8', 'ignore')


def _extract_doc_text(filename: str, blob: bytes):
    """上傳檔 → 純文字。支援 .txt/.md/.docx；.doc 不支援。回 (text, error)。"""
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    if ext in ('txt', 'md', 'markdown', 'text', ''):
        return _decode_text(blob), None
    if ext == 'docx':
        try:
            import io
            from docx import Document
            doc = Document(io.BytesIO(blob))
            parts = [p.text for p in doc.paragraphs]
            for tbl in doc.tables:               # 表格文字也納入
                for row in tbl.rows:
                    parts.append(' '.join(c.text for c in row.cells))
            return '\n'.join(parts), None
        except Exception as e:
            return None, f'.docx 解析失敗（{filename}）：{e}'
    if ext == 'doc':
        return None, f'「{filename}」是舊版 .doc，請在 Word 另存為 .docx 或 .txt 再上傳。'
    return None, f'不支援的檔案型別「{filename}」（支援 .txt / .md / .docx）。'
