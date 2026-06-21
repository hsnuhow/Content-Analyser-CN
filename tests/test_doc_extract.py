# -*- coding: utf-8 -*-
"""doc_extract characterization 測試（app/doc_extract.py）。

鎖住上傳檔文字抽取行為。.txt/.md/.doc/不支援 路徑純 stdlib、可本地測；
.docx 需 python-docx（正式環境有；本地若無則該例自動跳過）。

可直接執行：python3 tests/test_doc_extract.py　｜　相容 pytest
"""
import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "app"))

import doc_extract as d  # noqa: E402


def test_txt_decodes_utf8():
    txt, err = d._extract_doc_text('note.txt', '哈囉 world'.encode('utf-8'))
    assert err is None and txt == '哈囉 world'

def test_md_decodes():
    txt, err = d._extract_doc_text('a.md', b'# Heading')
    assert err is None and txt == '# Heading'

def test_no_extension_treated_as_text():
    txt, err = d._extract_doc_text('README', b'plain')
    assert err is None and txt == 'plain'

def test_bad_utf8_ignored_not_crash():
    txt, err = d._extract_doc_text('a.txt', b'ok\xff\xfeend')
    assert err is None and 'ok' in txt and 'end' in txt   # 壞位元組被忽略、不崩

def test_legacy_doc_rejected():
    txt, err = d._extract_doc_text('old.doc', b'x')
    assert txt is None and '舊版 .doc' in err

def test_unsupported_type():
    txt, err = d._extract_doc_text('sheet.xlsx', b'x')
    assert txt is None and '不支援的檔案型別' in err

def test_docx_if_available():
    try:
        import docx  # noqa: F401
    except Exception:
        return  # 本地無 python-docx → 跳過（正式環境有）
    import io
    from docx import Document
    doc = Document()
    doc.add_paragraph('段落一')
    doc.add_paragraph('段落二')
    buf = io.BytesIO()
    doc.save(buf)
    txt, err = d._extract_doc_text('x.docx', buf.getvalue())
    assert err is None and '段落一' in txt and '段落二' in txt


def _run():
    tests = [(n, f) for n, f in sorted(globals().items())
             if n.startswith("test_") and callable(f)]
    passed = failed = 0
    for name, fn in tests:
        try:
            fn(); passed += 1
        except AssertionError as e:
            failed += 1; print(f"  ✗ {name}  {e}")
        except Exception as e:
            failed += 1; print(f"  ✗ {name}  (例外) {e}")
    print(f"doc_extract：{passed} passed, {failed} failed（共 {len(tests)}）")
    return failed == 0


if __name__ == "__main__":
    sys.exit(0 if _run() else 1)
